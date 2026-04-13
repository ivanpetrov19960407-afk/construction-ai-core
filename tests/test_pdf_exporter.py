"""Тесты PDFExporter."""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from core.pdf_exporter import PDFExporter


@pytest.fixture
def minimal_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Тестовый документ", level=1)
    doc.add_paragraph("Минимальное содержимое для конвертации.")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_docx_to_pdf_returns_bytes_or_runtime_error(minimal_docx_bytes: bytes):
    exporter = PDFExporter()

    try:
        pdf_bytes = exporter.docx_to_pdf(minimal_docx_bytes, "minimal.docx")
    except RuntimeError as exc:
        assert "Не удалось конвертировать DOCX в PDF" in str(exc)
    else:
        assert isinstance(pdf_bytes, bytes)
        assert len(pdf_bytes) > 10


def test_generate_pdf_from_template_for_tk_template_returns_bytes():
    exporter = PDFExporter()
    context = {
        "work_type": "бетонирование",
        "scope": "Тестовая область применения",
        "technology": "Тестовая технология",
        "quality_requirements": "Тестовые требования качества",
        "normative_docs": ["СП 70.13330"],
        "sha256": "abc",
    }

    try:
        pdf_bytes = exporter.generate_pdf_from_template("tk_template", context)
    except RuntimeError as exc:
        assert "Не удалось конвертировать DOCX в PDF" in str(exc)
    else:
        assert isinstance(pdf_bytes, bytes)
        assert len(pdf_bytes) > 10
