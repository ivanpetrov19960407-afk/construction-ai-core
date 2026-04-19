"""Tests for RAG management endpoints."""

from __future__ import annotations

from fastapi.testclient import TestClient

from api.main import app
from api.routes import rag as rag_routes
from api.routes.auth import _create_token
from config.settings import settings


class _DummyCollection:
    def __init__(self, metadatas: list[dict]):
        self._metadatas = metadatas

    def get(self, include: list[str] | None = None) -> dict:
        return {"metadatas": self._metadatas}


class _DummyRAGEngine:
    def __init__(self):
        self.collection = _DummyCollection(
            [
                {"source": "sp_48.pdf"},
                {"source": "sp_48.pdf"},
                {"source": "gost_21.pdf"},
            ]
        )

    def ingest_pdf(self, filepath: str, source_name: str, metadata: dict | None = None) -> int:
        assert filepath.endswith(".pdf")
        assert source_name == "sp_48.pdf"
        if metadata is not None:
            assert metadata.get("session_id") == "chat-123"
        return 7

    def ingest_text(self, text: str, source_name: str, metadata: dict | None = None) -> int:
        assert source_name == "spec.docx"
        assert "Техническое задание" in text
        if metadata is not None:
            assert metadata.get("session_id") == "chat-docx"
        return 3


def test_rag_ingest_pdf_success(monkeypatch):
    old_api_keys = settings.api_keys
    old_admin_keys = settings.admin_api_keys
    settings.api_keys = ["admin-key"]
    settings.admin_api_keys = ["admin-key"]

    monkeypatch.setattr(rag_routes, "get_rag_engine", lambda: _DummyRAGEngine())

    parse_called = {"ok": False}

    def _fake_parse(file_bytes: bytes, filename: str):
        parse_called["ok"] = True
        assert filename == "sp_48.pdf"
        assert file_bytes.startswith(b"%PDF")
        return object()

    monkeypatch.setattr(rag_routes.pdf_parser, "parse", _fake_parse)

    try:
        with TestClient(app) as client:
            response = client.post(
                "/api/rag/ingest",
                files={
                    "file": ("sp_48.pdf", b"%PDF-1.4 fake", "application/pdf"),
                    "source_name": (None, "sp_48.pdf"),
                    "session_id": (None, "chat-123"),
                },
                headers={"X-API-Key": "admin-key"},
            )
    finally:
        settings.api_keys = old_api_keys
        settings.admin_api_keys = old_admin_keys

    assert response.status_code == 200
    assert response.json() == {"chunks_added": 7, "source": "sp_48.pdf"}
    assert parse_called["ok"] is True


def test_rag_ingest_docx_success(monkeypatch):
    from docx import Document

    old_api_keys = settings.api_keys
    old_admin_keys = settings.admin_api_keys
    settings.api_keys = ["admin-key"]
    settings.admin_api_keys = ["admin-key"]

    monkeypatch.setattr(rag_routes, "get_rag_engine", lambda: _DummyRAGEngine())

    doc = Document()
    doc.add_paragraph("Техническое задание")
    doc.add_paragraph("Проверка загрузки DOCX")

    import io

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    try:
        with TestClient(app) as client:
            response = client.post(
                "/api/rag/ingest",
                files={
                    "file": (
                        "spec.docx",
                        docx_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                    "source_name": (None, "spec.docx"),
                    "session_id": (None, "chat-docx"),
                },
                headers={"X-API-Key": "admin-key"},
            )
    finally:
        settings.api_keys = old_api_keys
        settings.admin_api_keys = old_admin_keys

    assert response.status_code == 200
    assert response.json() == {"chunks_added": 3, "source": "spec.docx"}


def test_rag_chat_upload_allows_non_admin_key(monkeypatch):
    old_api_keys = settings.api_keys
    old_admin_keys = settings.admin_api_keys
    settings.api_keys = ["user-key"]
    settings.admin_api_keys = []

    class _ChatUploadRAG:
        def __init__(self):
            self.collection = _DummyCollection([])

        def ingest_pdf(self, filepath: str, source_name: str, metadata: dict | None = None) -> int:
            assert filepath.endswith(".pdf")
            assert source_name == "chat-file.pdf"
            assert metadata == {"session_id": "chat-456"}
            return 5

    monkeypatch.setattr(rag_routes, "get_rag_engine", lambda: _ChatUploadRAG())
    monkeypatch.setattr(rag_routes.pdf_parser, "parse", lambda file_bytes, filename: object())

    try:
        with TestClient(app) as client:
            response = client.post(
                "/api/rag/chat-upload",
                files={
                    "file": ("chat-file.pdf", b"%PDF-1.4 fake", "application/pdf"),
                    "session_id": (None, "chat-456"),
                },
                headers={"X-API-Key": "user-key"},
            )
    finally:
        settings.api_keys = old_api_keys
        settings.admin_api_keys = old_admin_keys

    assert response.status_code == 200
    assert response.json() == {"chunks_added": 5, "source": "chat-file.pdf"}


def test_rag_ingest_docx_invalid_payload_returns_400(monkeypatch):
    old_api_keys = settings.api_keys
    old_admin_keys = settings.admin_api_keys
    settings.api_keys = ["admin-key"]
    settings.admin_api_keys = ["admin-key"]

    monkeypatch.setattr(rag_routes, "get_rag_engine", lambda: _DummyRAGEngine())

    try:
        with TestClient(app) as client:
            response = client.post(
                "/api/rag/ingest",
                files={
                    "file": ("broken.docx", b"not-a-real-docx", rag_routes.DOCX_MIME_TYPE),
                    "source_name": (None, "broken.docx"),
                },
                headers={"X-API-Key": "admin-key"},
            )
    finally:
        settings.api_keys = old_api_keys
        settings.admin_api_keys = old_admin_keys

    assert response.status_code == 400
    assert response.json() == {"detail": "Invalid DOCX file"}


def test_rag_sources_requires_admin_and_returns_counts(monkeypatch):
    old_api_keys = settings.api_keys
    old_admin_keys = settings.admin_api_keys
    settings.api_keys = ["valid-key", "admin-key"]
    settings.admin_api_keys = ["admin-key"]

    monkeypatch.setattr(rag_routes, "get_rag_engine", lambda: _DummyRAGEngine())

    try:
        with TestClient(app) as client:
            forbidden = client.get(
                "/api/rag/sources",
                headers={"X-API-Key": "valid-key"},
            )
            ok = client.get(
                "/api/rag/sources",
                headers={"X-API-Key": "admin-key"},
            )
    finally:
        settings.api_keys = old_api_keys
        settings.admin_api_keys = old_admin_keys

    assert forbidden.status_code == 403
    assert forbidden.json() == {"detail": "Admin role required"}

    assert ok.status_code == 200
    assert ok.json() == {
        "sources": [
            {"source": "gost_21.pdf", "chunks": 1},
            {"source": "sp_48.pdf", "chunks": 2},
        ]
    }


def test_rag_sources_allows_admin_jwt(monkeypatch):
    old_api_keys = settings.api_keys
    old_admin_keys = settings.admin_api_keys
    settings.api_keys = ["valid-key"]
    settings.admin_api_keys = []

    monkeypatch.setattr(rag_routes, "get_rag_engine", lambda: _DummyRAGEngine())

    admin_token = _create_token("admin-user", "admin")
    try:
        with TestClient(app) as client:
            response = client.get(
                "/api/rag/sources",
                headers={"Authorization": f"Bearer {admin_token}"},
            )
    finally:
        settings.api_keys = old_api_keys
        settings.admin_api_keys = old_admin_keys

    assert response.status_code == 200
    assert response.json() == {
        "sources": [
            {"source": "gost_21.pdf", "chunks": 1},
            {"source": "sp_48.pdf", "chunks": 2},
        ]
    }
