"""Generate endpoints — генерация документов."""

import asyncio
import json
import re
import tempfile
import uuid
from io import BytesIO
from datetime import datetime, timedelta, timezone
from enum import Enum
from pathlib import Path
from typing import Literal

import structlog
from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel, Field, field_validator
from sqlalchemy import create_engine, text

from agents.calculator import CalculatorAgent
from config.settings import settings
from core.billing import require_quota
from core.branding import BrandingConfig, get_branding
from core.cache import RedisCache
from core.export.onec_exporter import OneCExporter
from core.llm_router import LLMRouter
from core.multitenancy import get_tenant_id
from core.orchestrator import Orchestrator
from core.pdf_exporter import PDFExporter
from core.pdf_parser import PDFParser

router = APIRouter()
orchestrator = Orchestrator()
session_memory = orchestrator.session_memory
pdf_parser = PDFParser()
pdf_exporter = PDFExporter()
onec_exporter = OneCExporter()
redis_cache = RedisCache(settings.redis_url)
logger = structlog.get_logger("api.generate")

UTC = getattr(datetime, "UTC", timezone(timedelta(0)))

ALLOWED_UNITS = ["м³", "м²", "пог.м.", "шт.", "т", "кг"]
MAX_UPLOAD_SIZE_BYTES = 20 * 1024 * 1024
ANALYZE_TIMEOUT_SECONDS = 120
EXEC_ALBUM_SECTIONS = ("AR", "KZH", "KM", "KMD", "OV", "VK", "EM", "TX", "GP")
SESSION_TTL_SECONDS = 30 * 60
SESSION_CLEANUP_INTERVAL_SECONDS = 10 * 60
SESSION_STORE: dict[str, dict] = {}
_cleanup_task: asyncio.Task | None = None
_cleanup_lock = asyncio.Lock()


def _now_mono() -> float:
    """Текущее время event-loop в монотонной шкале."""
    return asyncio.get_event_loop().time()


def _is_session_expired(session_data: dict, now: float | None = None) -> bool:
    created_at = float(session_data.get("created_at", 0.0))
    current = now if now is not None else _now_mono()
    return (current - created_at) > SESSION_TTL_SECONDS


def _touch_session(
    session_id: str,
    *,
    text: str,
    doc_type: str,
    docx_bytes: bytes | None = None,
) -> None:
    SESSION_STORE[session_id] = {
        "text": text,
        "doc_type": doc_type,
        "docx_bytes": docx_bytes,
        "created_at": _now_mono(),
    }


def _get_live_session(session_id: str, expected_doc_type: str) -> dict | None:
    session_data = SESSION_STORE.get(session_id)
    if not isinstance(session_data, dict):
        return None
    if session_data.get("doc_type") != expected_doc_type:
        return None
    if _is_session_expired(session_data):
        SESSION_STORE.pop(session_id, None)
        return None
    return session_data


def _render_docx_from_text(text: str) -> bytes:
    """Сгенерировать минимальный DOCX-файл по тексту документа."""
    document = Document()
    for chunk in (text or "").splitlines():
        if chunk.strip():
            document.add_paragraph(chunk)
        else:
            document.add_paragraph("")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def cleanup_expired_sessions() -> None:
    """Фоновая очистка просроченных in-memory сессий."""
    while True:
        await asyncio.sleep(SESSION_CLEANUP_INTERVAL_SECONDS)
        now = _now_mono()
        stale_ids = [
            sid
            for sid, data in SESSION_STORE.items()
            if isinstance(data, dict) and _is_session_expired(data, now=now)
        ]
        for sid in stale_ids:
            SESSION_STORE.pop(sid, None)


async def _ensure_cleanup_task_started() -> None:
    global _cleanup_task  # noqa: PLW0603
    async with _cleanup_lock:
        if _cleanup_task is None or _cleanup_task.done():
            _cleanup_task = asyncio.create_task(cleanup_expired_sessions())


class TKRequest(BaseModel):
    """Запрос на генерацию технологической карты."""

    work_type: str = Field(min_length=5)
    object_name: str
    volume: float
    unit: str
    norms: list[str] = []
    role: str = "pto_engineer"
    session_id: str | None = None

    @field_validator("volume")
    @classmethod
    def validate_volume(cls, value: float) -> float:
        """Объём работ должен быть больше нуля."""
        if value <= 0:
            raise ValueError("volume must be > 0")
        return value

    @field_validator("unit")
    @classmethod
    def validate_unit(cls, value: str) -> str:
        """Единица измерения должна быть из разрешённого списка."""
        if value not in ALLOWED_UNITS:
            raise ValueError(f"unit must be one of: {', '.join(ALLOWED_UNITS)}")
        return value


class TKResponse(BaseModel):
    """Ответ на генерацию технологической карты."""

    result: str
    session_id: str
    document: dict
    agents_used: list[str]
    confidence: float | None
    sha256: str | None


class WorkItem(BaseModel):
    """Позиция работ для КС-2/КС-3."""

    name: str
    unit: str
    volume: float
    norm_hours: float
    price_per_unit: float


class KSRequest(BaseModel):
    """Запрос на генерацию КС-2/КС-3."""

    object_name: str
    contract_number: str
    period_from: str
    period_to: str
    work_items: list[WorkItem] = Field(min_length=1, max_length=50)
    role: str = "pto_engineer"
    session_id: str | None = None


class KSResponse(BaseModel):
    """Ответ на генерацию КС-2/КС-3."""

    result: str
    session_id: str
    ks2: dict
    ks3: dict
    docx_bytes_key: str
    total_cost: float
    total_hours: float
    sha256: str | None


class EstimateWorkItem(BaseModel):
    """Позиция работ для сметного расчёта по ГЭСН/ТСН."""

    work_type: str
    volume: float
    unit: str


class EstimateRequest(BaseModel):
    """Запрос на сметный расчёт по расценкам."""

    work_items: list[EstimateWorkItem] = Field(min_length=1, max_length=100)
    region: str = "Москва"


class EstimateResponse(BaseModel):
    """Ответ сметного расчёта."""

    items: list[dict]
    total_cost: float
    total_labor_hours: float
    region: str
    indexed_total_cost: float
    docx_available: bool = False


class LetterType(str, Enum):  # noqa: UP042
    """Тип делового письма."""

    REQUEST = "запрос"
    CLAIM = "претензия"
    NOTIFICATION = "уведомление"
    RESPONSE = "ответ"


class LetterRequest(BaseModel):
    """Запрос на генерацию делового письма."""

    letter_type: LetterType
    addressee: str
    subject: str
    body_points: list[str] = Field(min_length=1, max_length=10)
    contract_number: str | None = None
    include_npa: bool = True
    role: str = "foreman"
    session_id: str | None = None


class LetterResponse(BaseModel):
    """Ответ на генерацию делового письма."""

    result: str
    session_id: str
    document: dict
    agents_used: list[str]
    legal_references: list[str]
    confidence: float | None


class PPRRequest(BaseModel):
    """Запрос на генерацию ППР."""

    work_type: str
    object_name: str
    developer: str = "ИИ-помощник Construction AI"
    start_date: str
    duration_days: int = Field(gt=0)
    workers_count: int = Field(gt=0)
    role: str = "pto_engineer"
    session_id: str | None = None
    export_format: Literal["docx", "pdf"] = "docx"
    norms: list[str] = []


class AlbumRequest(BaseModel):
    """Запрос на сборку исполнительного альбома."""

    project_id: str
    section: Literal["AR", "KZH", "KM", "KMD", "OV", "VK", "EM", "TX", "GP"]


def _fetch_approved_exec_docs(project_id: str, section: str, org_id: str = "default") -> list[dict]:
    """Получить утвержденные АОСР по проекту и разделу."""
    engine = create_engine(settings.database_url, future=True)
    query = text(
        """
        SELECT id, pdf_url, created_at
        FROM executive_docs
        WHERE project_id = :project_id
          AND org_id = :org_id
          AND discipline_section = :section
          AND status = 'approved'
        ORDER BY created_at ASC
        """
    )
    with engine.connect() as conn:
        rows = (
            conn.execute(
                query,
                {"project_id": project_id, "section": section, "org_id": org_id},
            )
            .mappings()
            .all()
        )
    return [dict(row) for row in rows]


def _render_exec_album_pdf(
    project_id: str,
    section: str,
    docs: list[dict],
    branding: BrandingConfig,
) -> bytes:
    """Сформировать PDF альбома через WeasyPrint."""
    from weasyprint import HTML

    template_path = Path("templates/exec_album_cover.html")
    cover_template = template_path.read_text(encoding="utf-8")
    generated_at = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
    docs_html = "\n".join(
        (
            f"<li><strong>Документ #{index}</strong>: "
            f'<a href="{doc["pdf_url"]}">{doc["pdf_url"]}</a>'
            f" <span>(created_at: {doc.get('created_at', '-')})</span></li>"
        )
        for index, doc in enumerate(docs, start=1)
    )
    html = (
        cover_template.replace("{{ project_id }}", project_id)
        .replace("{{ section }}", section)
        .replace("{{ generated_at }}", generated_at)
        .replace("{{ doc_count }}", str(len(docs)))
        .replace("{{ docs_list }}", docs_html)
        .replace("{{ company_name }}", branding.company_name)
        .replace("{{ logo_url }}", branding.logo_url)
    )
    return HTML(string=html, base_url=str(Path.cwd())).write_pdf()


def _upload_album_bytes(project_id: str, section: str, pdf_bytes: bytes) -> str:
    """Загрузить PDF альбома в S3/MinIO и вернуть presigned URL."""
    try:
        import boto3
    except ImportError as exc:  # pragma: no cover - runtime dependency
        raise HTTPException(
            status_code=500,
            detail="boto3 is required for S3/MinIO upload",
        ) from exc

    timestamp = datetime.now(UTC).strftime("%Y%m%d%H%M%S")
    object_key = f"{project_id}/{section}_{timestamp}.pdf"

    client = boto3.client(
        "s3",
        endpoint_url=settings.s3_endpoint_url or None,
        aws_access_key_id=settings.s3_access_key or None,
        aws_secret_access_key=settings.s3_secret_key or None,
        region_name=settings.s3_region,
        use_ssl=settings.s3_use_ssl,
    )
    client.put_object(
        Bucket=settings.s3_bucket_albums,
        Key=object_key,
        Body=pdf_bytes,
        ContentType="application/pdf",
    )
    return str(
        client.generate_presigned_url(
            "get_object",
            Params={"Bucket": settings.s3_bucket_albums, "Key": object_key},
            ExpiresIn=3600,
        )
    )


def _upsert_generated_doc(
    doc_id: str,
    doc_type: str,
    payload: dict,
    org_id: str = "default",
) -> None:
    """Сохранить структурированные данные документа в generated_docs."""
    engine = create_engine(settings.database_url, future=True)
    create_table_query = text(
        """
        CREATE TABLE IF NOT EXISTS generated_docs (
            id TEXT NOT NULL,
            type TEXT NOT NULL,
            org_id TEXT NOT NULL DEFAULT 'default',
            payload TEXT NOT NULL,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            PRIMARY KEY (id, type, org_id)
        )
        """
    )
    upsert_query = text(
        """
        INSERT INTO generated_docs (id, type, org_id, payload, created_at, updated_at)
        VALUES (:id, :type, :org_id, :payload, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
        ON CONFLICT(id, type, org_id) DO UPDATE SET
            payload = excluded.payload,
            updated_at = CURRENT_TIMESTAMP
        """
    )
    with engine.begin() as conn:
        conn.execute(create_table_query)
        conn.execute(
            upsert_query,
            {
                "id": doc_id,
                "type": doc_type,
                "org_id": org_id,
                "payload": json.dumps(payload, ensure_ascii=False),
            },
        )


@router.post(
    "/generate/exec-album",
    summary="Сборка исполнительного альбома",
    description="Собирает PDF-альбом по утвержденным АОСР выбранного раздела.",
)
async def generate_exec_album(
    payload: AlbumRequest,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
    _quota: None = Depends(require_quota("exec_albums")),
):
    """Собрать исполнительный альбом по проекту и разделу."""
    _ = (request, _quota)
    if payload.section not in EXEC_ALBUM_SECTIONS:
        raise HTTPException(status_code=422, detail="Invalid section")

    tenant = org_id or "default"
    try:
        docs = await asyncio.to_thread(
            _fetch_approved_exec_docs,
            project_id=payload.project_id,
            section=payload.section,
            org_id=tenant,
        )
    except TypeError:
        docs = await asyncio.to_thread(
            _fetch_approved_exec_docs,
            project_id=payload.project_id,
            section=payload.section,
        )
    if not docs:
        raise HTTPException(status_code=404, detail="Approved executive docs not found")

    branding = await get_branding(tenant)
    pdf_bytes = await asyncio.to_thread(
        _render_exec_album_pdf,
        project_id=payload.project_id,
        section=payload.section,
        docs=docs,
        branding=branding,
    )
    presigned_url = await asyncio.to_thread(
        _upload_album_bytes,
        project_id=payload.project_id,
        section=payload.section,
        pdf_bytes=pdf_bytes,
    )
    return {
        "url": presigned_url,
        "doc_count": len(docs),
        "section": payload.section,
    }


@router.post(
    "/generate/tk",
    response_model=TKResponse,
    summary="Генерация технологической карты",
    description="Генерирует ТК на основе вида работ, объёмов и нормативов.",
    openapi_extra={
        "requestBody": {
            "required": True,
            "content": {
                "application/json": {
                    "example": {
                        "work_type": "Устройство монолитной плиты",
                        "object_name": "ЖК Северный квартал, корпус 3",
                        "volume": 120.5,
                        "unit": "м³",
                        "norms": ["СП 70.13330", "СП 48.13330"],
                        "role": "pto_engineer",
                        "session_id": "1c53f75f-4036-4b8f-9046-a46ad9175d1f",
                    }
                }
            },
        }
    },
)
async def generate_tk(
    payload: TKRequest,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
):
    """Генерация технологической карты (ТК) через orchestrator."""
    _ = (request, org_id)
    await _ensure_cleanup_task_started()
    session_id = payload.session_id or str(uuid.uuid4())
    norms_text = ", ".join(payload.norms) if payload.norms else "не указаны"
    message = (
        "Сформируй технологическую карту на русском языке.\n"
        f"Вид работ: {payload.work_type}.\n"
        f"Наименование объекта: {payload.object_name}.\n"
        f"Объём работ: {payload.volume} {payload.unit}.\n"
        f"Нормативы: {norms_text}.\n"
        "Подготовь структурированный документ для последующего DOCX-форматирования."
    )

    try:
        result = await orchestrator.process(
            message=message,
            session_id=session_id,
            role=payload.role,
            intent="generate_tk",
            extra_state={
                "work_type": payload.work_type,
                "object_name": payload.object_name,
                "volume": payload.volume,
                "unit": payload.unit,
                "norms": payload.norms,
                "tk_generator_input": {
                    "work_type": payload.work_type,
                    "object_name": payload.object_name,
                    "volume": payload.volume,
                    "unit": payload.unit,
                    "norms": payload.norms,
                },
            },
        )
    except Exception as exc:
        logger.exception("generate_tk_llm_error", session_id=session_id, error=str(exc))
        raise HTTPException(
            status_code=503,
            detail="LLM временно недоступен, попробуйте позже",
        ) from exc

    state = result.get("state", {}) if isinstance(result, dict) else {}
    docx_bytes = state.get("docx_bytes")
    verification = state.get("verification", {}) if isinstance(state, dict) else {}
    sha256 = (verification or {}).get("sha256")
    if isinstance(docx_bytes, bytes):
        await session_memory.save_document(
            session_id=session_id,
            doc_type="tk",
            filename=f"tk_{session_id}.docx",
            docx_bytes=docx_bytes,
            sha256=sha256,
        )

    document = (
        state.get("final_output") or state.get("docx_payload") or {"content": result.get("reply")}
    )
    tk_bridge_result = result.get("tk_bridge_result") if isinstance(result, dict) else None
    if isinstance(document, dict) and tk_bridge_result:
        document["tk_generator"] = tk_bridge_result
    result_text = str(result.get("reply") or "")
    if not result_text and isinstance(document, dict):
        result_text = json.dumps(document, ensure_ascii=False, indent=2)
    _touch_session(
        session_id,
        text=result_text,
        doc_type="tk",
        docx_bytes=docx_bytes if isinstance(docx_bytes, bytes) else None,
    )

    return TKResponse(
        result=result_text,
        session_id=result["session_id"],
        document=document,
        agents_used=result.get("agents_used", []),
        confidence=result.get("confidence"),
        sha256=sha256,
    )


LEGAL_REFERENCE_PATTERN = re.compile(
    r"(ст\.\s*\d+(?:\.\d+)?\s*[А-ЯA-Zа-яa-zЁё\s]{0,25}РФ|ФЗ-\d+)",
    flags=re.IGNORECASE,
)


def _extract_legal_references(history: list[dict]) -> list[str]:
    """Извлечь ссылки на НПА из history записей Legal Expert."""
    references: list[str] = []
    for item in history:
        if not isinstance(item, dict):
            continue
        agent_name = str(item.get("agent_name", "")).lower().replace(" ", "")
        if agent_name != "legalexpert":
            continue

        output = str(item.get("output", ""))
        matches = LEGAL_REFERENCE_PATTERN.findall(output)
        for match in matches:
            normalized = " ".join(match.strip().split())
            if normalized and normalized not in references:
                references.append(normalized)
    return references


@router.post(
    "/generate/letter",
    response_model=LetterResponse,
    summary="Генерация делового письма",
    description="Формирует деловое письмо по заданному типу, адресату, теме и ключевым тезисам.",
    openapi_extra={
        "requestBody": {
            "required": True,
            "content": {
                "application/json": {
                    "example": {
                        "letter_type": "запрос",
                        "addressee": "ООО СтройМонтаж",
                        "subject": "Предоставление исполнительной документации",
                        "body_points": [
                            "Просим направить акты скрытых работ",
                            "Указать сроки предоставления",
                        ],
                        "contract_number": "Д-17/2026",
                        "include_npa": True,
                        "role": "foreman",
                        "session_id": "802d28ad-2381-4837-97d6-4096bb5659fd",
                    }
                }
            },
        }
    },
)
async def generate_letter_v2(
    payload: LetterRequest,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
):
    """Генерация делового письма через orchestrator."""
    _ = (request, org_id)
    await _ensure_cleanup_task_started()
    session_id = payload.session_id or str(uuid.uuid4())
    body_points_text = "; ".join(payload.body_points)
    contract_number = payload.contract_number or "не указан"
    message = (
        f"Тип: {payload.letter_type.value}. "
        f"Получатель: {payload.addressee}. "
        f"Тема: {payload.subject}. "
        f"Тезисы: {body_points_text}. "
        f"Договор: {contract_number}."
    )

    try:
        result = await orchestrator.process(
            message=message,
            session_id=session_id,
            role=payload.role,
            intent="generate_letter",
            include_legal_expert=payload.include_npa,
        )
    except Exception as exc:
        logger.exception("generate_letter_llm_error", session_id=session_id, error=str(exc))
        raise HTTPException(
            status_code=503,
            detail="LLM временно недоступен, попробуйте позже",
        ) from exc

    state = result.get("state", {}) if isinstance(result, dict) else {}
    history = state.get("history", []) if isinstance(state, dict) else []
    document = state.get("docx_payload") or {"content": result.get("reply")}
    legal_references = _extract_legal_references(history)
    verification = state.get("verification", {}) if isinstance(state, dict) else {}
    sha256 = (verification or {}).get("sha256")
    docx_bytes = state.get("docx_bytes")
    if isinstance(docx_bytes, bytes):
        await session_memory.save_document(
            session_id=session_id,
            doc_type="letter",
            filename=f"letter_{session_id}.docx",
            docx_bytes=docx_bytes,
            sha256=sha256,
        )
    result_text = str(result.get("reply") or "")
    if not result_text and isinstance(document, dict):
        result_text = json.dumps(document, ensure_ascii=False, indent=2)
    _touch_session(
        session_id,
        text=result_text,
        doc_type="letter",
        docx_bytes=docx_bytes if isinstance(docx_bytes, bytes) else None,
    )

    return LetterResponse(
        result=result_text,
        session_id=result["session_id"],
        document=document,
        agents_used=result.get("agents_used", []),
        legal_references=legal_references,
        confidence=result.get("confidence"),
    )


@router.post(
    "/generate/ppr",
    summary="Генерация ППР",
    description=(
        "Формирует проект производства работ (ППР) с возможностью дальнейшей выгрузки в DOCX/PDF."
    ),
    openapi_extra={
        "requestBody": {
            "required": True,
            "content": {
                "application/json": {
                    "example": {
                        "work_type": "Монтаж металлоконструкций",
                        "object_name": "Складской комплекс А-12",
                        "developer": "ПТО ООО Генподряд",
                        "start_date": "2026-05-10",
                        "duration_days": 45,
                        "workers_count": 18,
                        "role": "pto_engineer",
                        "session_id": "ea8f6325-26a2-4bd1-a831-c4b8ec3c5aa5",
                        "export_format": "docx",
                    }
                }
            },
        }
    },
)
async def generate_ppr(
    payload: PPRRequest,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
):
    """Генерация ППР в DOCX/PDF."""
    _ = (request, org_id)
    session_id = payload.session_id or str(uuid.uuid4())
    message = (
        "Сформируй проект производства работ (ППР) на русском языке. "
        f"Вид работ: {payload.work_type}. "
        f"Объект: {payload.object_name}. "
        f"Дата начала: {payload.start_date}. "
        f"Длительность: {payload.duration_days} дней. "
        f"Численность: {payload.workers_count} чел."
    )
    if len(payload.norms) > 5:
        queued = await redis_cache.enqueue(
            "doc_generation",
            {
                "task_type": "generate_ppr",
                "session_id": session_id,
                "payload": payload.model_dump(),
            },
        )
        if not queued:
            logger.warning(
                "ppr_enqueue_failed",
                session_id=session_id,
                queue="doc_generation",
            )
            raise HTTPException(
                status_code=503,
                detail="Task queue is unavailable. Please retry later.",
            )

        return {
            "session_id": session_id,
            "status": "queued",
            "queue": "doc_generation",
            "message": "Тяжелая генерация ППР поставлена в очередь.",
        }

    context = {
        "work_type": payload.work_type,
        "object_name": payload.object_name,
        "general_data": (
            f"Объект {payload.object_name}. Начало работ: {payload.start_date}. "
            "Продолжительность: "
            f"{payload.duration_days} дней. "
            f"Состав звена: {payload.workers_count} чел."
        ),
        "ppr_sections": [
            {"name": "Общие данные", "description": "Сведения об объекте и исходных данных"},
            {"name": "Календарный план", "description": "Очередность и сроки выполнения работ"},
            {"name": "Охрана труда", "description": "Меры безопасности и контроль рисков"},
        ],
        "site_plan_description": "Размещение временных зданий, складов и подъездных путей.",
        "schedule_table": [
            {
                "stage": "Подготовительный этап",
                "start_date": payload.start_date,
                "duration_days": 5,
            },
            {
                "stage": "Основные работы",
                "start_date": payload.start_date,
                "duration_days": payload.duration_days,
            },
        ],
        "safety_measures": [
            "Проведение вводного и целевого инструктажа.",
            "Применение СИЗ и ограждений опасных зон.",
            "Ежесменный контроль состояния техники.",
        ],
        "normative_docs": "СП 48.13330, СП 70.13330, ТК РФ ст. 214",
        "developer": payload.developer,
        "start_date": payload.start_date,
    }
    try:
        result = await orchestrator.process(
            message=message,
            session_id=session_id,
            role=payload.role,
            intent="generate_tk",
            extra_state={
                "template_name": "ppr_template",
                "template_context": context,
                "export_format": payload.export_format,
            },
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"LLM processing error: {exc}") from exc

    state = result.get("state", {}) if isinstance(result, dict) else {}
    docx_bytes = state.get("docx_bytes")
    if not isinstance(docx_bytes, bytes):
        raise HTTPException(status_code=500, detail="Generated DOCX is empty")
    await session_memory.save_document(
        session_id=session_id,
        doc_type="ppr",
        filename=f"ppr_{session_id}.docx",
        docx_bytes=docx_bytes,
        sha256=(state.get("verification", {}) or {}).get("sha256"),
    )

    return {
        "session_id": result["session_id"],
        "document": state.get("final_output", {}),
        "agents_used": result.get("agents_used", []),
        "confidence": result.get("confidence"),
    }


@router.post(
    "/generate/ks",
    response_model=KSResponse,
    summary="Генерация КС-2 и КС-3",
    description="Генерирует формы КС-2/КС-3 по перечню выполненных работ за указанный период.",
    openapi_extra={
        "requestBody": {
            "required": True,
            "content": {
                "application/json": {
                    "example": {
                        "object_name": "БЦ Речной, секция Б",
                        "contract_number": "К-2026-11",
                        "period_from": "2026-03-01",
                        "period_to": "2026-03-31",
                        "work_items": [
                            {
                                "name": "Устройство бетонной подготовки",
                                "unit": "м³",
                                "volume": 80.0,
                                "norm_hours": 12.5,
                                "price_per_unit": 5200.0,
                            }
                        ],
                        "role": "pto_engineer",
                        "session_id": "7042f67a-7df6-4cf4-8fd4-065a2b3fac58",
                    }
                }
            },
        }
    },
)
async def generate_ks(
    payload: KSRequest,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
):
    """Генерация КС-2/КС-3 через orchestrator pipeline."""
    _ = request
    await _ensure_cleanup_task_started()
    tenant = org_id or "default"
    session_id = payload.session_id or str(uuid.uuid4())
    work_names = ", ".join(item.name for item in payload.work_items)
    message = (
        "Сформируй КС-2/КС-3 на русском языке. "
        f"Объект: {payload.object_name}. "
        f"Договор: {payload.contract_number}. "
        f"Период: {payload.period_from} — {payload.period_to}. "
        f"Наименования работ: {work_names}."
    )

    extra_state = {
        "object_name": payload.object_name,
        "contract_number": payload.contract_number,
        "period_from": payload.period_from,
        "period_to": payload.period_to,
        "calculation_params": {"work_items": [item.model_dump() for item in payload.work_items]},
        "context": (f"Подготовь описательную часть КС-2 для следующих работ: {work_names}."),
    }

    try:
        result = await orchestrator.process(
            message=message,
            session_id=session_id,
            role=payload.role,
            intent="generate_ks",
            extra_state=extra_state,
        )
    except Exception as exc:
        logger.exception("generate_ks_llm_error", session_id=session_id, error=str(exc))
        raise HTTPException(
            status_code=503,
            detail="LLM временно недоступен, попробуйте позже",
        ) from exc

    state = result.get("state", {}) if isinstance(result, dict) else {}
    ks2 = state.get("ks2_data", {})
    ks3 = state.get("ks3_data", {})
    total_cost = float(ks2.get("total_cost", 0.0))
    total_hours = float(ks2.get("total_hours", 0.0))

    docx_bytes = state.get("docx_bytes")
    docx_bytes_key = session_id
    verification = state.get("verification", {}) if isinstance(state, dict) else {}
    sha256 = (verification or {}).get("sha256")
    if isinstance(docx_bytes, bytes):
        await session_memory.save_document(
            session_id=docx_bytes_key,
            doc_type="ks",
            filename=f"ks_{docx_bytes_key}.docx",
            docx_bytes=docx_bytes,
            sha256=sha256,
        )
    result_text = str(result.get("reply") or "")
    if not result_text:
        result_text = json.dumps(
            {"ks2": ks2 if isinstance(ks2, dict) else {}, "ks3": ks3 if isinstance(ks3, dict) else {}},
            ensure_ascii=False,
            indent=2,
        )
    _touch_session(
        session_id,
        text=result_text,
        doc_type="ks",
        docx_bytes=docx_bytes if isinstance(docx_bytes, bytes) else None,
    )
    try:
        await asyncio.to_thread(
            _upsert_generated_doc,
            doc_id=docx_bytes_key,
            doc_type="ks2",
            payload=ks2 if isinstance(ks2, dict) else {},
            org_id=tenant,
        )
    except TypeError:
        await asyncio.to_thread(
            _upsert_generated_doc,
            doc_id=docx_bytes_key,
            doc_type="ks2",
            payload=ks2 if isinstance(ks2, dict) else {},
        )

    return KSResponse(
        result=result_text,
        session_id=result["session_id"],
        ks2=ks2,
        ks3=ks3,
        docx_bytes_key=docx_bytes_key,
        total_cost=total_cost,
        total_hours=total_hours,
        sha256=sha256,
    )


@router.post(
    "/generate/estimate",
    response_model=EstimateResponse,
    summary="Сметный расчёт по ТСН/ГЭСН",
    description="Выполняет ориентировочный расчёт стоимости и трудозатрат по каталогу расценок.",
)
async def generate_estimate(
    payload: EstimateRequest,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
):
    """Сметный калькулятор по справочнику расценок с региональным индексом."""
    _ = (request, org_id)
    calculator = CalculatorAgent(LLMRouter())
    estimate = calculator._calculate_estimate([item.model_dump() for item in payload.work_items])
    indexed_total_cost = calculator._apply_index(
        base_cost=float(estimate["total_cost"]),
        region=payload.region,
    )

    return EstimateResponse(
        items=estimate["items"],
        total_cost=float(estimate["total_cost"]),
        total_labor_hours=float(estimate["total_labor_hours"]),
        region=payload.region,
        indexed_total_cost=indexed_total_cost,
        docx_available=False,
    )


def _is_tender_document(filename: str, text_chunks: list[str]) -> bool:
    """Определить, относится ли документ к тендерным."""
    indicators = ("тендер", "закупк", "конкурс", "аукцион", "44-фз", "223-фз")
    haystack = f"{filename} {' '.join(text_chunks[:4])}".lower()
    return any(indicator in haystack for indicator in indicators)


def _extract_risks(analysis: str) -> list[str]:
    """Выделить строки с рисками из итогового анализа."""
    risks: list[str] = []
    for line in analysis.splitlines():
        normalized = line.strip(" -•\t")
        if normalized and "риск" in normalized.lower():
            risks.append(normalized)
    return risks


@router.post(
    "/analyze/document",
    summary="Анализ загруженного документа",
    description=(
        "Анализирует PDF-документ и возвращает "
        "структурированный результат с рисками и рекомендациями."
    ),
    openapi_extra={
        "requestBody": {
            "required": True,
            "content": {
                "multipart/form-data": {
                    "schema": {
                        "type": "object",
                        "properties": {
                            "file": {"type": "string", "format": "binary"},
                            "role": {"type": "string", "default": "tender_specialist"},
                            "session_id": {"type": "string"},
                        },
                        "required": ["file"],
                    },
                    "example": {
                        "role": "tender_specialist",
                        "session_id": "0b55b82d-d086-4f45-81d6-627dfd6d9fd7",
                    },
                }
            },
        }
    },
)
async def analyze_document(
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
    file: UploadFile = File(...),
    role: str = Form("tender_specialist"),
    session_id: str | None = Form(None),
):
    """Анализ загруженного PDF-документа через orchestrator."""
    _ = (request, org_id)
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Supported format: PDF only")

    file_bytes = await file.read()
    if len(file_bytes) > MAX_UPLOAD_SIZE_BYTES:
        raise HTTPException(status_code=413, detail="File size exceeds 20 MB limit")

    try:
        parsed = pdf_parser.parse(file_bytes=file_bytes, filename=file.filename)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"PDF parsing error: {exc}") from exc

    normative_refs = pdf_parser.extract_normative_refs("\n".join(parsed.text_chunks))
    document_intent = (
        "analyze_tender" if _is_tender_document(parsed.filename, parsed.text_chunks) else "chat"
    )
    session_id = session_id or str(uuid.uuid4())
    message = (
        f"Проанализируй документ: {parsed.filename}.\n"
        f"Текстовые фрагменты:\n{chr(10).join(parsed.text_chunks)}\n"
        f"Нормативные ссылки: {', '.join(normative_refs) if normative_refs else 'не найдены'}."
    )

    try:
        result = await asyncio.wait_for(
            orchestrator.process(
                message=message,
                session_id=session_id,
                role=role,
                intent=document_intent,
            ),
            timeout=ANALYZE_TIMEOUT_SECONDS,
        )
    except TimeoutError as exc:
        raise HTTPException(status_code=504, detail="Processing timeout (120 seconds)") from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"LLM processing error: {exc}") from exc

    analysis = str(result.get("reply") or "")
    return {
        "analysis": analysis,
        "normative_refs": normative_refs,
        "risks": _extract_risks(analysis),
        "session_id": result.get("session_id", session_id),
    }


@router.get("/generate/tk/{session_id}/download")
async def download_tk_docx(
    session_id: str,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
):
    """Скачать ранее сгенерированный DOCX по session_id."""
    _ = (request, org_id)
    await _ensure_cleanup_task_started()
    tk_session = _get_live_session(session_id, expected_doc_type="tk")
    if not tk_session:
        raise HTTPException(status_code=404, detail="Сессия не найдена или истекла")

    docx_bytes = tk_session.get("docx_bytes")
    if not isinstance(docx_bytes, bytes):
        docx_bytes = _render_docx_from_text(str(tk_session.get("text", "")))
        tk_session["docx_bytes"] = docx_bytes

    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = Path(tmp_file.name)
    with tmp_file:
        tmp_file.write(docx_bytes)

    return FileResponse(
        path=tmp_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"tk-{session_id}.docx",
    )


@router.get("/generate/letter/{session_id}/download")
async def download_letter_docx(
    session_id: str,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
):
    """Скачать ранее сгенерированный DOCX письма по session_id."""
    _ = (request, org_id)
    await _ensure_cleanup_task_started()
    letter_session = _get_live_session(session_id, expected_doc_type="letter")
    if not letter_session:
        raise HTTPException(status_code=404, detail="Сессия не найдена или истекла")

    docx_bytes = letter_session.get("docx_bytes")
    if not isinstance(docx_bytes, bytes):
        docx_bytes = _render_docx_from_text(str(letter_session.get("text", "")))
        letter_session["docx_bytes"] = docx_bytes

    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = Path(tmp_file.name)
    with tmp_file:
        tmp_file.write(docx_bytes)

    return FileResponse(
        path=tmp_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"letter-{session_id}.docx",
    )


@router.get("/generate/ks/{session_id}/download")
async def download_ks_docx(
    session_id: str,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
):
    """Скачать ранее сгенерированный DOCX КС-2/КС-3 по session_id."""
    _ = (request, org_id)
    await _ensure_cleanup_task_started()
    ks_session = _get_live_session(session_id, expected_doc_type="ks")
    if not ks_session:
        raise HTTPException(status_code=404, detail="Сессия не найдена или истекла")

    docx_bytes = ks_session.get("docx_bytes")
    if not isinstance(docx_bytes, bytes):
        docx_bytes = _render_docx_from_text(str(ks_session.get("text", "")))
        ks_session["docx_bytes"] = docx_bytes

    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = Path(tmp_file.name)
    with tmp_file:
        tmp_file.write(docx_bytes)

    return FileResponse(
        path=tmp_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"ks-{session_id}.docx",
    )


@router.get("/generate/ks2/{doc_id}/1c-xml")
async def export_ks2_1c_xml(
    doc_id: str,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
):
    """Экспорт КС-2 в XML-формат для импорта в 1С."""
    _ = request
    tenant = org_id or "default"
    try:
        xml_bytes = await onec_exporter.export_ks2_to_xml(doc_id=doc_id, org_id=tenant)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc

    return Response(
        content=xml_bytes,
        media_type="application/xml",
        headers={"Content-Disposition": f'attachment; filename="ks2_{doc_id}.xml"'},
    )


@router.get("/generate/m29/{project_id}/1c-xml")
async def export_m29_1c_xml(
    project_id: str,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
    period: str = Query(..., description="Период в формате YYYY-MM"),
):
    """Экспорт М-29 в XML-формат для импорта в 1С."""
    _ = (request, org_id)
    try:
        xml_bytes = await onec_exporter.export_m29_to_xml(project_id=project_id, period=period)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    return Response(
        content=xml_bytes,
        media_type="application/xml",
        headers={"Content-Disposition": f'attachment; filename="m29_{project_id}_{period}.xml"'},
    )


@router.get("/generate/letter/{session_id}/download")
async def download_letter_docx(
    session_id: str,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
    format: Literal["docx", "pdf"] = Query(default="docx"),
):
    """Скачать ранее сгенерированный DOCX/PDF письма по session_id."""
    _ = (request, org_id)
    documents = await session_memory.get_session_documents(session_id)
    letter_document = next((doc for doc in documents if doc.get("doc_type") == "letter"), None)
    docx_bytes = letter_document.get("docx_bytes") if letter_document else None
    if not docx_bytes:
        raise HTTPException(status_code=404, detail="DOCX not found for this session")

    if format == "pdf":
        docx_bytes = pdf_exporter.docx_to_pdf(docx_bytes, f"letter_{session_id}.docx")
    suffix = ".pdf" if format == "pdf" else ".docx"
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp_path = Path(tmp_file.name)
    with tmp_file:
        tmp_file.write(docx_bytes)

    return FileResponse(
        path=tmp_path,
        media_type=(
            "application/pdf"
            if format == "pdf"
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        filename=(
            f"letter_{session_id}.pdf"
            if format == "pdf"
            else (
                letter_document.get("filename") if letter_document else f"letter_{session_id}.docx"
            )
        ),
    )


@router.get("/generate/ppr/{session_id}/download")
async def download_ppr_docx(
    session_id: str,
    request: Request,
    org_id: str | None = Depends(get_tenant_id),
    format: Literal["docx", "pdf"] = Query(default="docx"),
):
    """Скачать ранее сгенерированный DOCX/PDF ППР по session_id."""
    _ = (request, org_id)
    documents = await session_memory.get_session_documents(session_id)
    ppr_document = next((doc for doc in documents if doc.get("doc_type") == "ppr"), None)
    docx_bytes = ppr_document.get("docx_bytes") if ppr_document else None
    if not docx_bytes:
        raise HTTPException(status_code=404, detail="DOCX not found for this session")

    if format == "pdf":
        docx_bytes = pdf_exporter.docx_to_pdf(docx_bytes, f"ppr_{session_id}.docx")
    suffix = ".pdf" if format == "pdf" else ".docx"
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp_path = Path(tmp_file.name)
    with tmp_file:
        tmp_file.write(docx_bytes)

    return FileResponse(
        path=tmp_path,
        media_type=(
            "application/pdf"
            if format == "pdf"
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        filename=(
            f"ppr_{session_id}.pdf"
            if format == "pdf"
            else (ppr_document.get("filename") if ppr_document else f"ppr_{session_id}.docx")
        ),
    )
