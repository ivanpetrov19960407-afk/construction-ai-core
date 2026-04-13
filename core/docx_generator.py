"""Утилита генерации DOCX из Jinja2-шаблонов."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate


class DocxGenerator:
    """Генератор DOCX на базе docxtpl."""

    def __init__(self, templates_dir: str | Path | None = None) -> None:
        base_dir = Path(__file__).resolve().parent.parent
        self.templates_dir = Path(templates_dir) if templates_dir else base_dir / "templates"

    def _create_default_tk_template(self, template_path: Path) -> None:
        """Создать минимальный tk_template.docx, если бинарный шаблон отсутствует."""
        template_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading("ТЕХНОЛОГИЧЕСКАЯ КАРТА на {{work_type}}", level=1)

        doc.add_heading("1. Область применения", level=2)
        doc.add_paragraph("{{scope}}")

        doc.add_heading("2. Организация и технология производства работ", level=2)
        doc.add_paragraph("{{technology}}")

        doc.add_heading("3. Требования к качеству работ", level=2)
        doc.add_paragraph("{{quality_requirements}}")

        doc.add_heading("4. Нормативные документы", level=2)
        doc.add_paragraph(
            "{% for doc_name in normative_docs %}"
            "• {{doc_name}}"
            "{% if not loop.last %}\n{% endif %}"
            "{% endfor %}"
        )

        footer = doc.sections[0].footer
        footer.paragraphs[0].text = "ГОСТ Р 21.1101 · SHA256: {{sha256}}"
        doc.save(template_path)

    def _create_default_ks_template(self, template_path: Path) -> None:
        """Создать минимальный ks_template.docx, если бинарный шаблон отсутствует."""
        template_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading("АКТ о приёмке выполненных работ (КС-2)", level=1)
        doc.add_paragraph("Объект: {{object_name}}")
        doc.add_paragraph("Договор: {{contract_number}}")
        doc.add_paragraph("Период: {{period_from}} — {{period_to}}")

        table = doc.add_table(rows=2, cols=6)
        header = table.rows[0].cells
        header[0].text = "№"
        header[1].text = "Наименование работ"
        header[2].text = "Ед.изм."
        header[3].text = "Объём"
        header[4].text = "Цена/ед."
        header[5].text = "Стоимость"

        row = table.rows[1].cells
        row[0].text = "{%tr for item in work_items %}{{item.index}}"
        row[1].text = "{{item.name}}"
        row[2].text = "{{item.unit}}"
        row[3].text = "{{item.volume}}"
        row[4].text = "{{item.price_per_unit}}"
        row[5].text = "{{item.subtotal_cost}}{%tr endfor %}"

        doc.add_paragraph("Итого: {{total_cost}} руб., {{total_hours}} чел.-ч")
        doc.save(template_path)

    def _ensure_template(self, template_name: str) -> Path:
        template_path = self.templates_dir / f"{template_name}.docx"
        if template_path.exists():
            return template_path

        if template_name == "tk_template":
            self._create_default_tk_template(template_path)
            return template_path
        if template_name == "ks_template":
            self._create_default_ks_template(template_path)
            return template_path

        raise FileNotFoundError(f"Template not found: {template_path}")

    def generate(self, template_name: str, context: dict) -> bytes:
        """Сгенерировать DOCX по шаблону и вернуть bytes."""
        template_path = self._ensure_template(template_name)

        tpl = DocxTemplate(str(template_path))
        tpl.render(context)

        buffer = BytesIO()
        tpl.save(buffer)
        return buffer.getvalue()

    def list_templates(self) -> list[str]:
        """Список доступных DOCX-шаблонов (без расширения)."""
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        templates = {
            path.stem for path in self.templates_dir.glob("*.docx") if path.is_file()
        }
        templates.add("tk_template")
        templates.add("ks_template")
        return sorted(templates)
