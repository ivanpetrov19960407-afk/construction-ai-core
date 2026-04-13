"""Экспорт DOCX в PDF с несколькими стратегиями."""

from __future__ import annotations

import html
import tempfile
from pathlib import Path

from docx import Document

from core.docx_generator import DocxGenerator


class PDFExporter:
    """Конвертер DOCX в PDF с fallback-стратегиями."""

    def _convert_with_docx2pdf(self, docx_path: Path, pdf_path: Path) -> bool:
        try:
            from docx2pdf import convert
        except ImportError:
            return False

        try:
            convert(str(docx_path), str(pdf_path))
        except Exception:
            return False

        return pdf_path.exists() and pdf_path.stat().st_size > 0

    def _docx_to_html(self, docx_path: Path) -> str:
        doc = Document(str(docx_path))
        blocks: list[str] = [
            "<html><head><meta charset='utf-8'>",
            "<style>body{font-family:DejaVu Sans,Arial,sans-serif;font-size:12pt;}"
            "h1,h2,h3{margin:12px 0 6px;}"
            "p{margin:0 0 8px;}"
            "table{border-collapse:collapse;width:100%;margin:8px 0;}"
            "td,th{border:1px solid #555;padding:6px;vertical-align:top;}"
            "</style></head><body>",
        ]

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            blocks.append(f"<p>{html.escape(text)}</p>")

        for table in doc.tables:
            blocks.append("<table>")
            for row in table.rows:
                blocks.append("<tr>")
                for cell in row.cells:
                    blocks.append(f"<td>{html.escape(cell.text)}</td>")
                blocks.append("</tr>")
            blocks.append("</table>")

        blocks.append("</body></html>")
        return "".join(blocks)

    def _convert_with_weasyprint(self, docx_path: Path, pdf_path: Path) -> bool:
        try:
            from weasyprint import HTML
        except Exception:
            return False

        html_content = self._docx_to_html(docx_path)
        try:
            HTML(string=html_content).write_pdf(str(pdf_path))
        except Exception:
            return False

        return pdf_path.exists() and pdf_path.stat().st_size > 0

    def docx_to_pdf(self, docx_bytes: bytes, filename: str) -> bytes:
        """Конвертирует DOCX → PDF.

        Стратегия:
        1. Попытка через docx2pdf (требует LibreOffice/Word)
        2. Fallback: конвертация через python-docx → HTML → weasyprint
        3. Если оба недоступны → RuntimeError
        """
        with tempfile.TemporaryDirectory() as tmp_dir:
            base_name = Path(filename).stem or "document"
            docx_path = Path(tmp_dir) / f"{base_name}.docx"
            pdf_path = Path(tmp_dir) / f"{base_name}.pdf"
            docx_path.write_bytes(docx_bytes)

            if self._convert_with_docx2pdf(docx_path, pdf_path):
                return pdf_path.read_bytes()

            if self._convert_with_weasyprint(docx_path, pdf_path):
                return pdf_path.read_bytes()

        raise RuntimeError(
            "Не удалось конвертировать DOCX в PDF: установите LibreOffice/Word для docx2pdf "
            "или установите зависимости weasyprint."
        )

    def generate_pdf_from_template(self, template_name: str, context: dict) -> bytes:
        """DOCX-шаблон → bytes DOCX → bytes PDF."""
        docx_bytes = DocxGenerator().generate(template_name, context)
        return self.docx_to_pdf(docx_bytes, template_name)
