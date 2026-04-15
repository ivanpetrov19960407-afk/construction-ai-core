"""Сравнение версий документов (TXT/DOCX) и формирование отчёта diff."""

from __future__ import annotations

from difflib import SequenceMatcher, unified_diff
from pathlib import Path

from docx import Document


class DocumentDiff:
    """Сервис сравнения текстовых версий документов."""

    SAFETY_KEYWORDS = (
        "безопас",
        "охран",
        "техника безопасности",
        "средства индивидуальной защиты",
        "сиз",
        "пожар",
        "эвакуац",
    )

    def compare_texts(self, text_v1: str, text_v2: str) -> dict:
        """Сравнить две текстовые версии и вернуть структурированный diff."""
        lines_v1 = text_v1.splitlines()
        lines_v2 = text_v2.splitlines()

        diff_lines = list(
            unified_diff(
                lines_v1,
                lines_v2,
                fromfile="v1",
                tofile="v2",
                lineterm="",
            )
        )

        added = [line[1:] for line in diff_lines if line.startswith("+") and not line.startswith("+++")]
        removed = [line[1:] for line in diff_lines if line.startswith("-") and not line.startswith("---")]
        changed_sections = [line for line in diff_lines if line.startswith("@@")]

        similarity_pct = round(SequenceMatcher(a=text_v1, b=text_v2).ratio() * 100, 2)

        critical_changes = [
            line
            for line in [*added, *removed]
            if any(keyword in line.lower() for keyword in self.SAFETY_KEYWORDS)
        ]
        if not critical_changes and (added or removed):
            safety_context_present = any(
                keyword in f"{text_v1}\n{text_v2}".lower() for keyword in self.SAFETY_KEYWORDS
            )
            if safety_context_present:
                critical_changes = [*removed, *added]

        return {
            "added": added,
            "removed": removed,
            "changed_sections": changed_sections,
            "similarity_pct": similarity_pct,
            "critical_changes": critical_changes,
        }

    def compare_docx(self, path_v1: str, path_v2: str) -> dict:
        """Прочитать DOCX-файлы и сравнить их как текст."""
        doc1 = Document(path_v1)
        doc2 = Document(path_v2)

        text_v1 = "\n".join(p.text for p in doc1.paragraphs if p.text.strip())
        text_v2 = "\n".join(p.text for p in doc2.paragraphs if p.text.strip())

        return self.compare_texts(text_v1=text_v1, text_v2=text_v2)

    def generate_diff_report(self, diff: dict) -> str:
        """Сформировать человеко-читаемый отчёт по diff на русском."""
        added_count = len(diff.get("added", []))
        removed_count = len(diff.get("removed", []))
        similarity_pct = float(diff.get("similarity_pct", 0.0))

        lines = [
            f"+ добавлено {added_count} строк",
            f"- удалено {removed_count} строк",
            f"Сходство версий: {similarity_pct:.2f}%",
        ]

        critical_changes = diff.get("critical_changes", [])
        if critical_changes:
            lines.append("Критические изменения (разделы по безопасности):")
            for line in critical_changes[:10]:
                lines.append(f"- {line}")
        else:
            lines.append("Критические изменения по безопасности не обнаружены.")

        changed_sections = diff.get("changed_sections", [])
        if changed_sections:
            lines.append("Изменённые секции:")
            lines.extend(changed_sections[:10])

        return "\n".join(lines)


def read_text_file(path: str) -> str:
    """Прочитать текстовый файл (utf-8) с fallback cp1251."""
    file_path = Path(path)
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return file_path.read_text(encoding="cp1251")
