"""Генерация DOCX-шаблонов в каталоге templates/."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = BASE_DIR / "templates"


def _apply_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def build_tk_template(output_path: Path) -> None:
    doc = Document()
    _apply_default_font(doc)

    _add_heading(doc, "ТЕХНОЛОГИЧЕСКАЯ КАРТА на {{work_type}}")
    doc.add_paragraph()

    sections = [
        ("Область применения", "{{scope}}"),
        ("Организация работ", "{{technology}}"),
        ("Технология производства", "{{technology}}"),
        ("Требования к качеству", "{{quality_requirements}}"),
        ("Техника безопасности", "{{safety}}"),
        ("Нормативные документы", "{{normative_docs}}"),
    ]

    for title, body in sections:
        _add_heading(doc, title)
        doc.add_paragraph(body)

    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "ГОСТ Р 21.1101 · SHA256: {{sha256}}"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def build_letter_template(output_path: Path) -> None:
    doc = Document()
    _apply_default_font(doc)

    header = doc.add_table(rows=1, cols=2)
    header.rows[0].cells[0].text = "{{sender_org}}"
    header.rows[0].cells[1].text = "{{addressee}}"
    header.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("Исх. №{{outgoing_number}} от {{date}}")
    doc.add_paragraph("Этап 24 — DOCX-шаблоны для всех типов документов (2 дня)")

    _add_heading(doc, "Тема: {{subject}}")
    doc.add_paragraph("{{body}}")
    doc.add_paragraph("Подпись: {{signer_position}} _________ {{signer_name}}")
    doc.add_paragraph("Правовое основание: {{legal_references}}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def build_ks_template(output_path: Path) -> None:
    doc = Document()
    _apply_default_font(doc)

    _add_heading(doc, "АКТ о приёмке выполненных работ (форма КС-2)")
    doc.add_paragraph("Объект: {{object_name}}, Договор: {{contract_number}}")
    doc.add_paragraph("Период: {{period_from}} — {{period_to}}")

    table = doc.add_table(rows=2, cols=6)
    headers = ["№", "Наименование", "Ед.изм.", "Объём", "Цена/ед.", "Стоимость"]
    for index, text in enumerate(headers):
        table.rows[0].cells[index].text = text

    row = table.rows[1].cells
    row[0].text = "{%tr for item in work_items %}{{item.index}}"
    row[1].text = "{{item.name}}"
    row[2].text = "{{item.unit}}"
    row[3].text = "{{item.volume}}"
    row[4].text = "{{item.price_per_unit}}"
    row[5].text = "{{item.subtotal_cost}}{%tr endfor %}"

    doc.add_paragraph("Итого: {{total_cost}} руб., {{total_hours}} чел.-ч")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def build_ppr_template(output_path: Path) -> None:
    doc = Document()
    _apply_default_font(doc)

    _add_heading(doc, "ПРОЕКТ ПРОИЗВОДСТВА РАБОТ на {{work_type}}")

    _add_heading(doc, "Общие данные")
    doc.add_paragraph("{{general_data}}")

    _add_heading(doc, "Состав ППР")
    sections_table = doc.add_table(rows=2, cols=2)
    sections_table.rows[0].cells[0].text = "Раздел"
    sections_table.rows[0].cells[1].text = "Описание"
    sections_table.rows[1].cells[0].text = "{%tr for row in ppr_sections %}{{row.name}}"
    sections_table.rows[1].cells[1].text = "{{row.description}}{%tr endfor %}"

    _add_heading(doc, "Стройгенплан")
    doc.add_paragraph("{{site_plan_description}}")

    _add_heading(doc, "Календарный план")
    schedule_table = doc.add_table(rows=2, cols=3)
    schedule_table.rows[0].cells[0].text = "Этап"
    schedule_table.rows[0].cells[1].text = "Дата начала"
    schedule_table.rows[0].cells[2].text = "Длительность, дни"
    schedule_table.rows[1].cells[0].text = "{%tr for item in schedule_table %}{{item.stage}}"
    schedule_table.rows[1].cells[1].text = "{{item.start_date}}"
    schedule_table.rows[1].cells[2].text = "{{item.duration_days}}{%tr endfor %}"

    _add_heading(doc, "Охрана труда")
    doc.add_paragraph(
        "{% for measure in safety_measures %}"
        "• {{measure}}"
        "{% if not loop.last %}\n{% endif %}"
        "{% endfor %}"
    )

    _add_heading(doc, "Нормативные документы")
    doc.add_paragraph("{{normative_docs}}")

    doc.add_paragraph("Разработал: {{developer}}")
    doc.add_paragraph("Дата: {{start_date}}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def generate_all_templates(templates_dir: Path = TEMPLATES_DIR) -> list[Path]:
    templates_dir.mkdir(parents=True, exist_ok=True)

    files = {
        "tk_template.docx": build_tk_template,
        "letter_template.docx": build_letter_template,
        "ks_template.docx": build_ks_template,
        "ppr_template.docx": build_ppr_template,
    }

    generated: list[Path] = []
    for filename, builder in files.items():
        output_path = templates_dir / filename
        builder(output_path)
        generated.append(output_path)

    return generated


if __name__ == "__main__":
    target_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else TEMPLATES_DIR
    created = generate_all_templates(target_dir)
    for path in created:
        print(f"generated: {path}")
