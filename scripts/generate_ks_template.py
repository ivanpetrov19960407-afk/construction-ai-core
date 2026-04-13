"""Генерация шаблона templates/ks_template.docx через python-docx."""

from pathlib import Path

from docx import Document


def build_ks_template(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("АКТ о приёмке выполненных работ (КС-2)", level=1)
    doc.add_paragraph("Объект: {{object_name}}")
    doc.add_paragraph("Договор: {{contract_number}}")
    doc.add_paragraph("Период: {{period_from}} — {{period_to}}")

    table = doc.add_table(rows=2, cols=6)
    headers = ["№", "Наименование работ", "Ед.изм.", "Объём", "Цена/ед.", "Стоимость"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    row = table.rows[1].cells
    row[0].text = "{%tr for item in work_items %}{{item.index}}"
    row[1].text = "{{item.name}}"
    row[2].text = "{{item.unit}}"
    row[3].text = "{{item.volume}}"
    row[4].text = "{{item.price_per_unit}}"
    row[5].text = "{{item.subtotal_cost}}{%tr endfor %}"

    doc.add_paragraph("Итого: {{total_cost}} руб., {{total_hours}} чел.-ч")
    doc.save(output_path)


if __name__ == "__main__":
    build_ks_template(Path("templates/ks_template.docx"))
    print("templates/ks_template.docx generated")
